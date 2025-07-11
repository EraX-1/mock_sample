import os

import pypdf as PyPDF2
from azure.ai.formrecognizer import FormRecognizerClient
from azure.core.credentials import AzureKeyCredential
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_community.document_loaders import AzureAIDocumentIntelligenceLoader

from src.config.azure_config import MOCK_CONFIG


class MockDocumentIntelligenceLoader:
    """ローカル開発用のDocument Intelligenceモックローダー"""

    def __init__(self, file_path: str, **kwargs):
        self.file_path = file_path

    def load(self) -> list[Document]:
        """ローカルファイルを解析してマークダウンに変換"""
        try:
            content = self._extract_text_from_file(self.file_path)
            # マークダウン形式に変換
            markdown_content = self._convert_to_markdown(content)

            return [
                Document(
                    page_content=markdown_content,
                    metadata={
                        "source": self.file_path,
                        "title": os.path.basename(self.file_path),
                        "mock_service": True,
                    },
                )
            ]
        except Exception as e:
            return [
                Document(
                    page_content=f"ファイルの読み込みに失敗しました: {str(e)}",
                    metadata={
                        "source": self.file_path,
                        "error": str(e),
                        "mock_service": True,
                    },
                )
            ]

    def _extract_text_from_file(self, file_path: str) -> str:
        """ファイルからテキストを抽出"""
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == ".pdf":
            return self._extract_from_pdf(file_path)
        elif file_extension in [".docx", ".doc"]:
            return self._extract_from_docx(file_path)
        elif file_extension == ".txt":
            with open(file_path, encoding="utf-8") as f:
                return f.read()
        else:
            return f"サポートされていないファイル形式です: {file_extension}"

    def _extract_from_pdf(self, file_path: str) -> str:
        """PDFからテキストを抽出"""
        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            return f"PDFの読み込みエラー: {str(e)}"

    def _extract_from_docx(self, file_path: str) -> str:
        """Word文書からテキストを抽出"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return f"Word文書の読み込みエラー: {str(e)}"

    def _convert_to_markdown(self, text: str) -> str:
        """テキストを簡易的なマークダウンに変換"""
        lines = text.split("\n")
        markdown_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                markdown_lines.append("")
                continue

            # 簡易的な見出し検出
            if len(line) < 50 and not line.endswith(".") and not line.endswith(","):
                markdown_lines.append(f"## {line}")
            else:
                markdown_lines.append(line)

        return "\n".join(markdown_lines)


class MockFormRecognizerClient:
    """ローカル開発用のForm Recognizerクライアントモック"""

    def __init__(self, **kwargs):
        pass

    def begin_recognize_content_from_url(self, url):
        """URLからコンテンツを認識（モック）"""
        return type(
            "MockResult",
            (),
            {
                "result": lambda: [
                    type(
                        "MockForm",
                        (),
                        {
                            "pages": [
                                type(
                                    "MockPage",
                                    (),
                                    {
                                        "lines": [
                                            type(
                                                "MockLine",
                                                (),
                                                {"text": "これはモック環境での文書解析結果です。"},
                                            )()
                                        ]
                                    },
                                )()
                            ]
                        },
                    )()
                ]
            },
        )()


class AzureAIDocumentIntelligence:
    def __init__(self):
        self.use_mock = MOCK_CONFIG["use_mock_services"]
        self.document_intelligence_key = os.environ.get(
            "DOCUMENT_INTELLIGENCE_KEY", "mock-doc-intel-key"
        )
        self.document_intelligence_endpoint = os.environ.get(
            "DOCUMENT_INTELLIGENCE_ENDPOINT",
            "https://mock.cognitiveservices.azure.com/",
        )

    def init_form_recognizer_client(self):
        if self.use_mock:
            return MockFormRecognizerClient()
        else:
            return FormRecognizerClient(
                endpoint=self.document_intelligence_endpoint,
                credential=AzureKeyCredential(self.document_intelligence_key),
            )

    def init_loader(self, file_path: str):
        if self.use_mock:
            return MockDocumentIntelligenceLoader(file_path=file_path)
        else:
            try:
                return AzureAIDocumentIntelligenceLoader(
                    file_path=file_path,
                    api_key=self.document_intelligence_key,
                    api_endpoint=self.document_intelligence_endpoint,
                    api_model="prebuilt-layout",
                    mode="markdown",
                )
            except Exception as e:
                # 詳細なエラー情報をログに出力
                print(f"AzureAIDocumentIntelligenceLoader初期化エラー: {e}")
                print(f"エンドポイント: {self.document_intelligence_endpoint}")
                print(f"APIキー: {self.document_intelligence_key[:10]}...")
                raise e
